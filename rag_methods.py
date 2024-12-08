import streamlit as st
import os
import dotenv
from langchain_community.vectorstores import Chroma
from langchain_openai import OpenAIEmbeddings
from langchain_core.prompts import ChatPromptTemplate, MessagesPlaceholder
from langchain.chains import create_history_aware_retriever, create_retrieval_chain
from langchain.chains.combine_documents import create_stuff_documents_chain
from time import time
import pandas as pd
from langchain.schema import Document
import spacy
import logging
from PyPDF2 import PdfReader
from docx import Document as DocxDocument


os.environ["USER_AGENT"] = "YourAppName/1.0"
# Load SpaCy model for semantic splitting
nlp = spacy.load("en_core_web_sm")

dotenv.load_dotenv()

DB_DOCS_LIMIT = 20  # Increase the document upload limit

def extract_text_by_page(file_path):
    """
    Extract text from a PDF file, page by page.
    """
    try:
        with open(file_path, 'rb') as pdf:
            reader = PdfReader(pdf)
            return [page.extract_text().strip() for page in reader.pages]
    except Exception as e:
        logging.error(f"Failed to extract text from PDF: {e}")
        return []

def extract_text_from_docx(file_path):
    """
    Extract text from a Word document, including headers, paragraphs, and tables.
    """
    try:
        doc = DocxDocument(file_path)
        content = []

        # Extract paragraphs
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if text:
                content.append(text)

        # Extract tables
        for table in doc.tables:
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_data:
                    content.append(" | ".join(row_data))  # Represent rows as pipe-separated

        return "\n\n".join(content)

    except Exception as e:
        logging.error(f"Error processing Word document {file_path}: {e}")
        return ""



def load_doc_to_db():
    """
    Handle document uploads. Use semantic chunking to split documents and store them in the vector database.
    """
    # Check if there are uploaded documents in the session state
    if "rag_docs" in st.session_state and st.session_state.rag_docs:
        docs = []  # List to store processed documents
        unique_sources = set(st.session_state.rag_sources)  # Track uploaded files to avoid duplicates

        # Iterate through the uploaded files
        for doc_file in st.session_state.rag_docs:
            # Check if the file is unique and within the upload limit
            if doc_file.name not in unique_sources and len(unique_sources) < DB_DOCS_LIMIT:
                # Ensure the directory for storing source files exists
                os.makedirs("source_files", exist_ok=True)
                file_path = f"./source_files/{doc_file.name}"

                # Save the uploaded file to the local directory
                with open(file_path, "wb") as file:
                    file.write(doc_file.read())

                try:
                    # Process PDF files
                    if doc_file.type == "application/pdf":
                        raw_pages = extract_text_by_page(file_path)  # Extract text from each page
                        raw_text = "\n\n".join(raw_pages)  # Combine text from all pages
                        docs.append(Document(page_content=raw_text))  # Create a Document object

                    # Process Word documents
                    elif doc_file.name.endswith(".docx"):
                        raw_text = extract_text_from_docx(file_path)  # Extract text from the Word document
                        if raw_text.strip():  # Check if content was extracted
                            docs.append(Document(page_content=raw_text))  # Add to docs
                        else:
                            st.warning(f"No content extracted from {doc_file.name}.")  # Warn if empty

                    # Process plain text and markdown files
                    elif doc_file.type in ["text/plain", "text/markdown"]:
                        with open(file_path, "r", encoding="utf-8") as file:
                            raw_text = file.read()  # Read the file content
                        docs.append(Document(page_content=raw_text))  # Add to docs

                    # Process Excel files
                    elif doc_file.name.endswith((".xls", ".xlsx")):
                        excel_data = pd.read_excel(file_path, engine="openpyxl")  # Load Excel data
                        json_data = excel_data.to_json(orient="records")  # Convert to JSON format
                        docs.append(Document(page_content=json_data))  # Add to docs

                    # Process CSV files
                    elif doc_file.name.endswith(".csv"):
                        csv_data = pd.read_csv(file_path, encoding="utf-8")  # Load CSV data
                        json_data = csv_data.to_json(orient="records")  # Convert to JSON format
                        docs.append(Document(page_content=json_data))  # Add to docs

                    # Handle unsupported file types
                    else:
                        st.warning(f"Document type {doc_file.type} not supported.")  # Warn user
                        continue  # Skip unsupported files

                    # Add the file name to the unique sources list
                    unique_sources.add(doc_file.name)
                    st.session_state.rag_sources = list(unique_sources)  # Update session state

                except Exception as e:
                    # Log an error if the document fails to process
                    st.error(f"Error processing {doc_file.name}: {e}")

        # If documents were successfully processed
        if docs:
            # Split and load documents into the vector database
            _split_and_load_docs(docs)
            st.toast("Documents loaded successfully.", icon="✅")  # Show success notification
        else:
            # Error message if the document limit is reached
            st.error(f"Maximum number of documents reached ({DB_DOCS_LIMIT}).")




def _split_and_load_docs(docs, max_tokens=800, overlap_tokens=200):
    """
    Split documents into token-based chunks with overlap and load them into the vector database.
    """
    chunks = []

    for doc in docs:
        # Approximate token count: 4 characters ≈ 1 token
        token_multiplier = 4

        # Break content into sentences
        doc_nlp = nlp(doc.page_content)
        sentences = [sentence.text for sentence in doc_nlp.sents]

        current_chunk = []
        current_chunk_size = 0

        for sentence in sentences:
            # Approximate the token count for the sentence
            sentence_size = len(sentence) // token_multiplier

            # Add sentence to the current chunk
            if current_chunk_size + sentence_size > max_tokens:
                # Save the current chunk if it's not empty
                if current_chunk:
                    chunks.append(" ".join(current_chunk))

                # Start the next chunk with the overlap
                overlap_start = max(0, len(current_chunk) - (overlap_tokens // token_multiplier))
                current_chunk = current_chunk[overlap_start:]
                current_chunk_size = sum(len(sent) // token_multiplier for sent in current_chunk)

            current_chunk.append(sentence)
            current_chunk_size += sentence_size

        # Add the remaining chunk
        if current_chunk:
            chunks.append(" ".join(current_chunk))

    # Filter out empty chunks
    chunks = [chunk for chunk in chunks if chunk.strip()]

    # Store chunks in session state for display
    if "chunked_knowledge" not in st.session_state:
        st.session_state.chunked_knowledge = []
    st.session_state.chunked_knowledge.extend(chunks)

    # Create Document objects for each chunk
    document_chunks = [Document(page_content=chunk) for chunk in chunks if chunk.strip()]

    # Check if vector_db is initialized
    if "vector_db" not in st.session_state or st.session_state.vector_db is None:
        st.session_state.vector_db = initialize_vector_db(document_chunks)
    else:
        st.session_state.vector_db.add_documents(document_chunks)
        st.session_state.vector_db.persist()  # Persist changes to database




def initialize_vector_db(docs):
    """
    Initialize a vector database and store documents with embeddings and metadata.
    """
    # Add metadata to each document
    for i, doc in enumerate(docs):
        doc.metadata = {"index": i}

    # Ensure database folder exists
    os.makedirs("./chroma_db", exist_ok=True)

    # Generate embeddings
    embeddings = OpenAIEmbeddings(api_key=st.session_state.openai_api_key)
    document_texts = [doc.page_content for doc in docs]
    try:
        embeddings_data = embeddings.embed_documents(document_texts)
        if not embeddings_data:
            raise ValueError("Generated embeddings are empty.")
    except Exception as e:
        logging.error(f"Error generating embeddings: {e}")
        raise

    # Initialize Chroma vector database
    vector_db = Chroma.from_documents(
        documents=docs,
        embedding=embeddings,
        persist_directory="./chroma_db",  # Persist directory for Chroma
        collection_name=f"{str(time()).replace('.', '')[:14]}_{st.session_state['session_id']}",
    )
    vector_db.persist()  # Ensure changes are saved

    return vector_db





def _get_context_retriever_chain(vector_db, llm):
    retriever = vector_db.as_retriever(
        search_type="similarity",
        search_kwargs={"k": 5}  # Retrieve fewer documents for relevance, 
                                # Relevance: Increase k if you want to broaden the scope of retrieved documents.
                                # Efficiency: Decrease k if performance or relevance sufficiency is a concern.
    )

    prompt = ChatPromptTemplate.from_messages([
        MessagesPlaceholder(variable_name="messages"),
        ("user", "{input}"),
        ("system", "Use the retrieved knowledge to craft a relevant response."),
    ])

    retriever_chain = create_history_aware_retriever(llm, retriever, prompt)
    return retriever_chain



def get_conversational_rag_chain(llm):
    if "vector_db" not in st.session_state or not st.session_state.vector_db:
        raise ValueError("No vector database found. Please upload a knowledge base or use the default LLM.")

    retriever_chain = _get_context_retriever_chain(st.session_state.vector_db, llm)
    prompt = ChatPromptTemplate.from_messages([
        ("system", """
        
# Role
You are a highly skilled and knowledgeable AI assistant named Zorro. Your primary role is to assist with Axent-related tasks, but you are also capable of answering general questions or helping with tasks like coding, troubleshooting, and providing explanations in a human-like tonality. Your expertise in Axent's internal processes and your ability to interpret historical faults to provide accurate solutions are crucial to the long-term success of the company.

{context}
# Task
1. When a user sends you a question or a message, always search through the knowledge base using RAG methods to see if there is anything relevant or related that can help the user with their question.
2. If you are able to retrieve the correct and useful data from the knowledge base, return a message to the user with the correct information as a short and brief summary, and ask the user if they would like more info on the certain topic.
3. If you are not able to find any relevant data within the knowledge base (such that the user could also be asking an unrelated question to Axent and their internal processes), then proceed to use normal Claude model functionality to help the user with any of their queries.
4. Kindly ask the user if they would like to ask any more questions or need further clarification.
5. Finally, before outputing your response, make sure that your response is unique and not simply copy pasted from the knowledge base, it's important to have unique answers that are different everytime, but still capture the same inherit meaning. Finally, do not repeat the same message twice but worded differently when retrieving from the knowledge base.
6. Never mention anything about the Axent knowledge base, so for example, if someone asks you a question, do not start the conversation by saying "Okay, let me see what I can find in the Axent knowledge base", or anything starting with "Based on...".

# Specifics
- The Axent knowledge base contains large amounts of data that relates to all of their internal processes. This can include simple questions about certain design topics, knowledge bases where you are able to interpret historical faults to see how they were fixed and recommend similar solutions, certain PCB repair data, etc.
- Your role as a support agent for Axent is crucial to the long-term success of the company, and it is extremely important that you are able to retrieve relevant information and, where applicable, provide recommendations or solutions as to how certain things can be fixed.
- When helping employees, use the following PCB repair flowchart to guide them to the right solution more quickly:

A[Visual Inspection  
Check for damaged components,  
broken tracks and signs of damage] --> B[Can the PCB be  
powered up?]  
B -->|No| C[Use test equipment  
e.g., Multimeter:
- Check power rails for shorts  
- Check fuses for open circuit  
- Check caps/inductors for shorts  
- Check resistors for open circuit  
- Measure resistor values]  
B -->|Yes| D[Power up PCB]  

C --> E
E --> F[Is a reference  
PCB available?]  

F -->|No| G[Check for design  
similarities]  
F -->|Yes| H[Check all components and ICs]  

G --> H  
H --> I[Replace components  
as required]  
I --> B  

D --> J[- Check current consumption  
- Use current limiter  
- Check PCB for heat  
- Use FLIR camera for heat spots]  
J --> K[Check all voltages]  

K --> L[- Measure all test points  
- Measure regulators, converters  
- Measure transformers  
- Measure Vcc on familiar ICs  
- Check power LED for correct  
colour]  

L --> M[Run custom tests]  

M --> N[- Check switches, LEDs, etc.  
- Check displays]  


- When providing solutions, focus on general guidance rather than overly specific details. For example, instead of "R319, C161 out of alignment, Reflowed u525," provide advice like "visually inspect all components for proper alignment and reflow as needed."
- **It's crucial that your responses are concise and to the point.** Avoid long paragraphs and aim for a maximum of 2-3 sentences per response. If the user needs more information, they can always ask follow-up questions.

# Context
Axent is a company that specialises in designing and manufacturing electronic controllers. Their products are used in a wide range of applications, from industrial automation to consumer electronics. As an AI assistant, your role is to support Axent's employees by providing them with accurate and timely information to help them troubleshoot issues, repair PCBs, and optimize their designs.
- The founder Geoff Fontaine worked at the local cricket centre and had to change the scoreboards manually, so he thought "how can I automate this", and then did just so from his garage.

The knowledge base you have access to contains a wealth of information on Axent's internal processes, design guidelines, and historical fault data. By leveraging this information, you can provide valuable insights and recommendations to employees, helping them work more efficiently and effectively.

Your ability to understand the context of each query and provide relevant, concise answers is essential to the success of Axent's operations. By assisting employees with their day-to-day tasks and helping them overcome challenges, you directly contribute to the company's growth and success.
         
# Notes
- If the query relates to Axent, prioritise the relevant Axent knowledge base.
- If the query is unrelated or the knowledge base doesn't contain relevant information, use your general AI capabilities to provide a thoughtful, accurate, and helpful response.
- Always aim to be concise and professional in your answers.
- Do not respond to the user by saying "According to the information provided," as it sounds unprofessional and not very human-like.
- Make sure responses do not use excess tokens if not necessary; answers should be straight to the point, with a maximum of 2-3 sentences."
- Do not start conversations by saying "According to the information in the knowledge base". This sounds unnatural and kills the user engagement.
- **Never** mention the specific name of the knowledge base file that you are retrieving information from (if relevant), as this comes off unnatural to the user. Make it seem as though you know everything naturally, and not explictely mentioning that you are retrieving the information from a certain named knowledge base, so, do not say "Based on [insert knowledge base]".
    """),  # Custom prompt, can modify for better compactness and token efficiency if needed.
        MessagesPlaceholder(variable_name="messages"),
        ("user", "{input}"),
    ])
    
    stuff_documents_chain = create_stuff_documents_chain(llm, prompt)

    return create_retrieval_chain(retriever_chain, stuff_documents_chain)



def stream_llm_rag_response(llm_stream, messages):
    conversation_rag_chain = get_conversational_rag_chain(llm_stream)
    response_message = ""
    
    # Try retrieving relevant results
    try:
        # Use the RAG retriever chain
        for chunk in conversation_rag_chain.pick("answer").stream({"messages": messages[:-1], "input": messages[-1].content}):
            response_message += chunk
            yield chunk
    except ValueError:
        # Fallback to default Claude response
        for chunk in llm_stream.stream(messages):
            response_message += chunk.content
            yield chunk

    # Return the response message
    return response_message



